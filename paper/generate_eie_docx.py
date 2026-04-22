#!/usr/bin/env python3
"""Generate Electrical Engineering & Electromechanics (NTU "KhPI") DOCX.

Target journal: Electrical Engineering & Electromechanics, NTU "KhPI"
  https://eie.khpi.edu.ua/
Scopus Q2 (2024), Category A (Ukraine), DOI prefix 10.20998
APC: 20 EUR per A4 page

Format spec:
  - Times New Roman 10pt body (captions 9pt)
  - Margins 2.5cm top/bottom/left, 1.5cm right
  - Abstract 250-300 words with 7 mandatory subsections:
      Introduction, Problem, Goal, Methodology, Results,
      Scientific novelty, Practical value
  - Keywords: 4-8 words
  - References: IEEE-style, 15-60 items
      70% DOI, 50% Scopus/WoS last 5 years
  - Formulas in Microsoft Equation (OMML)
  - Figures 300 dpi, BMP/JPG

Content source: paper/paper.tex (elsarticle LaTeX version).
Ukrainian translation is handled by EE&E Editorial Board —
the authors submit only English.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap

nsmap["m"] = "http://schemas.openxmlformats.org/officeDocument/2006/math"

FONT = "Times New Roman"
FONT_SIZE = Pt(10)
FONT_SIZE_TITLE = Pt(11)
FONT_SIZE_CAPTION = Pt(9)
FONT_SIZE_TABLE = Pt(9)
FONT_SIZE_SMALL = Pt(9)
FIRST_INDENT = 0.5  # cm


# ---------- formatting primitives ----------

def _force_font(run, size=FONT_SIZE, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:cs"), FONT)
    for attr in ["asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"]:
        rFonts.attrib.pop(qn(f"w:{attr}"), None)


def add_para(doc, text, *, size=FONT_SIZE, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0,
             indent=None, line=1.15):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _force_font(run, size, bold, italic)
    pf = p.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if indent is not None:
        pf.first_line_indent = Cm(indent)
    return p


def add_mixed(doc, parts, *, size=FONT_SIZE, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              before=0, after=0, indent=None, line=1.15):
    """parts: list of (text, bold, italic)."""
    p = doc.add_paragraph()
    for text, bold, italic in parts:
        run = p.add_run(text)
        _force_font(run, size, bold, italic)
    pf = p.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if indent is not None:
        pf.first_line_indent = Cm(indent)
    return p


def add_heading_numbered(doc, text):
    """'1. Introduction' style — bold, left-aligned per EE&E template example."""
    return add_para(doc, text, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
                    before=8, after=2, line=1.15)


def add_heading_unnumbered(doc, text):
    return add_para(doc, text, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
                    before=8, after=2, line=1.15)


def add_body(doc, text, *, indent=FIRST_INDENT):
    return add_para(doc, text, indent=indent, after=2)


def add_figure(doc, fig_num, caption, image_path, width_cm=8.5):
    """Single-column figure with caption. 8.5 cm fits A4 column."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(6)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0
    if os.path.exists(image_path):
        run = p.add_run()
        run.add_picture(image_path, width=Cm(width_cm))
    else:
        run = p.add_run(f"[Fig. {fig_num}: missing {os.path.basename(image_path)}]")
        _force_font(run, FONT_SIZE_CAPTION, italic=True)

    pc = doc.add_paragraph()
    r_lbl = pc.add_run(f"Fig. {fig_num}. ")
    _force_font(r_lbl, FONT_SIZE_CAPTION, bold=True)
    r_txt = pc.add_run(caption)
    _force_font(r_txt, FONT_SIZE_CAPTION)
    pc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.paragraph_format.space_after = Pt(6)
    pc.paragraph_format.line_spacing = 1.15


def add_table(doc, headers, rows, caption_num, caption_text, bold_rows=None):
    """Table with label 'Table X' left-aligned above, then italic caption below."""
    bold_rows = set(bold_rows or [])

    plbl = doc.add_paragraph()
    r = plbl.add_run(f"Table {caption_num}")
    _force_font(r, FONT_SIZE_CAPTION, bold=True)
    plbl.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    plbl.paragraph_format.space_before = Pt(6)
    plbl.paragraph_format.space_after = Pt(1)
    plbl.paragraph_format.line_spacing = 1.15

    pcap = doc.add_paragraph()
    r = pcap.add_run(caption_text)
    _force_font(r, FONT_SIZE_CAPTION, italic=True)
    pcap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pcap.paragraph_format.space_after = Pt(3)
    pcap.paragraph_format.line_spacing = 1.15

    n_cols = len(headers)
    n_rows = len(rows) + 1
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.style = "Table Grid"

    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        _force_font(run, FONT_SIZE_TABLE, bold=True)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0

    for i, row in enumerate(rows):
        is_bold = i in bold_rows
        for j, val in enumerate(row):
            cell = tbl.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            _force_font(run, FONT_SIZE_TABLE, bold=is_bold)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.0

    # trailing empty paragraph for spacing
    doc.add_paragraph().paragraph_format.line_spacing = 1.0


def add_reference(doc, num, text):
    """Numbered reference, IEEE style, Times New Roman 10 pt."""
    p = doc.add_paragraph()
    run = p.add_run(f"{num}. {text}")
    _force_font(run, FONT_SIZE)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.15
    pf.left_indent = Cm(0.75)
    pf.first_line_indent = Cm(-0.75)


# ---------- OMML formulas ----------

def _omml_r(text, *, italic=True, bold=False):
    r = OxmlElement("m:r")
    rPr = OxmlElement("m:rPr")
    sty = OxmlElement("m:sty")
    if italic and bold:
        sty.set(qn("m:val"), "bi")
    elif italic:
        sty.set(qn("m:val"), "i")
    elif bold:
        sty.set(qn("m:val"), "b")
    else:
        sty.set(qn("m:val"), "p")
    rPr.append(sty)
    r.append(rPr)
    t = OxmlElement("m:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    return r


def _omml_text_italic(text):
    return _omml_r(text, italic=True)


def _omml_sub(base, sub):
    sSub = OxmlElement("m:sSub")
    e = OxmlElement("m:e")
    if isinstance(base, str):
        e.append(_omml_r(base))
    else:
        e.append(base)
    sSub.append(e)
    s = OxmlElement("m:sub")
    if isinstance(sub, str):
        s.append(_omml_r(sub, italic=False))
    else:
        s.append(sub)
    sSub.append(s)
    return sSub


def _omml_sup(base, sup):
    sSup = OxmlElement("m:sSup")
    e = OxmlElement("m:e")
    if isinstance(base, str):
        e.append(_omml_r(base))
    else:
        e.append(base)
    sSup.append(e)
    s = OxmlElement("m:sup")
    if isinstance(sup, str):
        s.append(_omml_r(sup, italic=False))
    else:
        s.append(sup)
    sSup.append(s)
    return sSup


def add_formula(doc, omml_elements, number):
    """Centered OMML math with right-side number."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.15

    oMath = OxmlElement("m:oMath")
    for el in omml_elements:
        oMath.append(el)
    p._element.append(oMath)

    run = p.add_run(f"  ({number})")
    _force_font(run, FONT_SIZE)
    return p


def add_inline_formula_para(doc, text, number):
    """Simpler fallback: use Cambria Math text for formulas that are too
    complex to hand-code in OMML. Editorial team usually retypes these."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Cambria Math"
    r.font.size = FONT_SIZE
    r.font.italic = True
    r.font.color.rgb = RGBColor(0, 0, 0)
    num = p.add_run(f"  ({number})")
    _force_font(num, FONT_SIZE)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.15
    return p


# =================================================================
def generate():
    doc = Document()

    # Page setup per EE&E spec
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = FONT_SIZE
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    script_dir = os.path.dirname(os.path.abspath(__file__))
    fig_dir = os.path.join(script_dir, "..", "assets")

    # -----------------------------------------------------------------
    # UDC (assigned by editor; placeholder)
    # -----------------------------------------------------------------
    add_para(doc, "UDC 621.311:004.8", bold=True,
             align=WD_ALIGN_PARAGRAPH.LEFT, after=6)

    # -----------------------------------------------------------------
    # Authors (English) — names, italic affiliations below
    # -----------------------------------------------------------------
    add_para(doc, "D. Voitekh, A. Tymoshenko",
             size=Pt(11), bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=4)

    # -----------------------------------------------------------------
    # Title (English) — CAPS, bold, centered
    # -----------------------------------------------------------------
    add_para(doc,
             "BRIDGING THE MV/LV GAP: VIRTUAL SLACK NODES AND "
             "POSITIONAL ENCODINGS FOR GNN-BASED POWER FLOW ON "
             "RADIAL DISTRIBUTION NETWORKS",
             size=FONT_SIZE_TITLE, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=6)

    # -----------------------------------------------------------------
    # Structured abstract (250-300 words, 7 subsections)
    # -----------------------------------------------------------------
    abstract_parts = [
        ("Introduction. ", True, False),
        ("Graph neural networks (GNNs) offer 10\u2013100\u00d7 speedups over "
         "Newton\u2013Raphson (NR) solvers for AC power flow approximation, "
         "but their accuracy degrades on low-voltage (LV) distribution grids "
         "relative to medium-voltage (MV) grids. ", False, False),
        ("Problem. ", True, False),
        ("A preliminary benchmarking of four GNN architectures on ten "
         "SimBench distribution grids revealed a 41% MAE gap between LV and "
         "MV grids (p=0.010) attributable to the long graph diameter of "
         "radial LV feeders (up to 56 hops), which exceeds the receptive "
         "field of a standard 4-layer GNN. ", False, False),
        ("Goal. ", True, False),
        ("To close the MV/LV accuracy gap without abandoning the "
         "computational advantage of GNN surrogates. ", False, False),
        ("Methodology. ", True, False),
        ("Three complementary techniques are evaluated on ten SimBench grids "
         "(four MV, six LV) with 2,000 load scenarios each and three random "
         "seeds: (1) virtual slack edges that add bidirectional shortcut "
         "connections from the slack bus to every node, reducing the "
         "effective graph diameter to two hops; (2) random walk positional "
         "encodings with k=16; and (3) residual GraphSAGE with eight layers. ",
         False, False),
        ("Results. ", True, False),
        ("The combined model reduces mean voltage-magnitude MAE by 43.8% "
         "(Wilcoxon p=0.002) and drops the LV/MV ratio from 1.41\u00d7 to "
         "1.21\u00d7, rendering the gap statistically non-significant on "
         "SimBench (p=0.182). Virtual slack edges alone contribute a 34% "
         "MAE reduction. Baseline MAE is almost perfectly rank-correlated "
         "with graph diameter (Spearman \u03c1=+0.95, p<10\u207b\u2074). "
         "A physics-informed voltage-smoothness loss yields no measurable "
         "benefit and a scale analysis explains why. ", False, False),
        ("Scientific novelty. ", True, False),
        ("The slack bus is reinterpreted as a physically meaningful virtual "
         "node, and rank-correlation evidence links baseline error to "
         "topological properties diagnosable before training. ", False, False),
        ("Practical value. ", True, False),
        ("The combined model remains 2\u20136\u00d7 faster than NR while "
         "roughly halving LV-grid error, supporting Monte Carlo hosting-"
         "capacity and voltage-quality studies where LV accuracy dominates "
         "the computational budget. References 36, tables 5, figures 3.",
         False, False),
    ]
    add_mixed(doc, abstract_parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              after=4, indent=FIRST_INDENT)

    # -----------------------------------------------------------------
    # Keywords
    # -----------------------------------------------------------------
    add_mixed(doc, [
        ("Key words: ", True, True),
        ("distribution networks; graph neural networks; positional "
         "encodings; power flow; virtual nodes; machine learning.",
         False, True),
    ], after=10)

    # Ukrainian abstract placeholder — editorial board will translate
    add_para(doc,
             "[Ukrainian abstract will be provided by the Editorial Board "
             "per journal policy.]",
             italic=True, size=FONT_SIZE_SMALL,
             align=WD_ALIGN_PARAGRAPH.LEFT, after=10)

    # =================================================================
    # 1. INTRODUCTION
    # =================================================================
    add_heading_numbered(doc, "Introduction.")

    add_body(doc,
        "The rapid integration of distributed energy resources into "
        "distribution networks necessitates frequent power flow "
        "calculations for real-time monitoring, voltage regulation, and "
        "hosting capacity assessment [1]. While the Newton\u2013Raphson "
        "(NR) method remains the standard solver, its computational cost "
        "becomes prohibitive for large-scale Monte Carlo studies and "
        "online applications that require thousands of power flow "
        "solutions per minute [2].")

    add_body(doc,
        "Graph neural networks have emerged as a promising surrogate for "
        "power flow, exploiting the natural graph structure of electrical "
        "networks to learn the mapping from load injections to bus "
        "voltages [3, 4]; recent surveys document rapid uptake across "
        "power-system applications [33]. Prior studies have demonstrated "
        "4\u2013100\u00d7 speedups over NR while maintaining sub-percent "
        "accuracy [5, 6]. However, the application of GNNs to distribution "
        "networks \u2014 particularly low-voltage (LV) grids \u2014 has "
        "received less attention despite being the setting where "
        "computational savings matter most.")

    add_body(doc,
        "In a preliminary benchmarking study, four GNN architectures were "
        "compared on ten SimBench distribution grids and a systematic "
        "performance gap was identified: LV grids exhibited 41% higher "
        "mean absolute error (MAE) than MV grids (p=0.010, Mann\u2013"
        "Whitney U test). This gap was attributed to two structural "
        "properties of radial LV feeders: (1) large graph diameter (up to "
        "56 hops), which far exceeds the receptive field of a 4-layer GNN, "
        "and (2) high bridge fraction, meaning that most edges are "
        "topologically critical and cannot be bypassed through alternative "
        "paths.")

    add_body(doc,
        "This paper proposes three targeted solutions. (1) Virtual slack "
        "edges: bidirectional shortcut edges from the slack bus to every "
        "node, reducing the effective graph diameter to at most two hops. "
        "(2) Random walk positional encodings: node features based on "
        "k-step return probabilities, providing each node with a "
        "topology-aware fingerprint. (3) Residual GraphSAGE: an eight-"
        "layer architecture with pre-normalization residual connections, "
        "expanding the receptive field while mitigating oversmoothing.")

    add_body(doc,
        "Each component is evaluated individually and in combination across "
        "all ten grids. The combined model achieves a 43.8% mean MAE "
        "reduction (p=0.002) and, critically, reduces the LV/MV "
        "performance ratio from 1.41\u00d7 to 1.21\u00d7, rendering it "
        "statistically non-significant (p=0.182). A physics-informed "
        "voltage smoothness loss provides no measurable improvement; a "
        "negative result reported and analyzed.")

    # =================================================================
    # 2. LITERATURE REVIEW
    # =================================================================
    add_heading_numbered(doc, "Literature review.")

    add_body(doc,
        "GNN-based power flow approximation builds on graph convolutional "
        "networks [3]. Donon et al. [4] introduced a message-passing "
        "framework that minimizes Kirchhoff violation for AC power flow, "
        "while Lin et al. [5] proposed PowerFlowNet, achieving 4\u201348\u00d7 "
        "speedups via a combined message-passing and high-order GCN "
        "architecture. Ringsquandl et al. [6] showed that GNNs on power "
        "grids benefit from exceptionally deep architectures (up to 13 "
        "layers), unlike typical graph benchmarks. Lopez-Garcia and "
        "Dom\u00ednguez-Navarro [29] introduced a typed graph formulation "
        "that distinguishes buses and lines as heterogeneous node types. "
        "Beinert et al. [27] forecast node-level power flow at transmission "
        "grid nodes with GNN plus embedding multi-task learning. Hu et al. "
        "[25] proposed an adaptive graph deep learning scheme that handles "
        "varying topologies during operation, while Ghamizi et al. [10] "
        "recently extended GNN power flow to unbalanced three-phase "
        "distribution systems with PowerFlowMultiNet.")

    add_body(doc,
        "Physics-informed formulations have also appeared: PINCO [7] "
        "incorporates power balance constraints into GNN training for "
        "optimal power flow; Zhang et al. [8] proposed a physics-informed "
        "line graph neural network; Gao et al. [30] embed power-flow "
        "physics directly into the graph convolution kernel to improve "
        "robustness to uncertain injections and topology changes; their "
        "follow-up work [31] generalizes the idea to optimal power flow. "
        "Deihim et al. [32] use a GNN to produce warm-start estimates "
        "that accelerate interior-point solvers for AC-OPF. For state "
        "estimation, Ngo et al. [26] combined graph structure with "
        "physics-based regularization, and Madbhavi et al. [28] "
        "demonstrated scalability of GNN-based estimators on realistic "
        "distribution systems. The PowerGraph benchmark [9] provides "
        "standardized GNN evaluation datasets but focuses on transmission "
        "networks. Suri and Mangal [11] combined GraphSAGE with GRU cells "
        "for topology-aware temporal prediction. Okoyomon and Goebel [12] "
        "proposed a generalizability framework for GNN-based power flow "
        "on SimBench distribution grids, using positional encodings to "
        "improve cross-topology transfer, but did not address the MV/LV "
        "accuracy gap.")

    add_body(doc,
        "A parallel line of work uses graph-structured learning for "
        "control rather than forward prediction: Chen et al. [34] applied "
        "multi-agent graph reinforcement learning with a physical-"
        "assistance mechanism to fast voltage regulation in PV-rich "
        "distribution networks, and Pei et al. [35] studied topology-"
        "robust voltage control via multi-task reinforcement learning. "
        "Operational risk assessment under evolving grid conditions has "
        "also been approached with GNNs [36]. Most existing work, however, "
        "targets transmission networks or optimal power flow. Distribution-"
        "level power flow \u2014 where radial topology creates the most "
        "acute structural challenges for GNNs \u2014 remains comparatively "
        "underexplored.")

    add_body(doc,
        "Standard message-passing GNNs are limited by the Weisfeiler\u2013"
        "Leman isomorphism test and cannot distinguish certain graph "
        "structures [13]. Positional encodings (PEs) address this by "
        "augmenting node features with structural information. Dwivedi et "
        "al. [14] introduced Laplacian eigenvector PEs for graph "
        "transformers, while Li et al. [15] proposed distance encodings "
        "including random walk landing probabilities. These techniques "
        "have proven effective in molecular property prediction and "
        "combinatorial optimization but have not been applied to power "
        "flow.")

    add_body(doc,
        "The oversmoothing phenomenon \u2014 where node representations "
        "converge to a common value as GNN depth increases \u2014 is well "
        "documented [16]. Residual connections [17], jumping knowledge "
        "(JKNet) [18], and DropEdge [19] have been proposed as "
        "mitigations. For power flow on radial grids, where the graph "
        "diameter can reach 56 hops, the ability to build deep networks "
        "without oversmoothing is essential. The virtual node technique, "
        "introduced by Gilmer et al. [20] for molecular graphs and later "
        "adopted in the Open Graph Benchmark [21], adds a global node "
        "connected to all others, enabling long-range information "
        "propagation in a single message-passing step. While effective in "
        "molecular and social network domains, this technique has not "
        "been explored for power flow.")

    # =================================================================
    # 3. PROBLEM FORMULATION AND METHODOLOGY
    # =================================================================
    add_heading_numbered(doc, "Problem formulation and methodology.")

    add_body(doc,
        "AC power flow is formulated as a supervised node regression "
        "problem on a graph G=(V,E), where V represents buses and E "
        "represents branches (lines and transformers). Each node i has "
        "input features x_i comprising active and reactive power "
        "injections (P_i, Q_i) and bus type indicators. The target "
        "outputs are voltage magnitude V_{m,i} and voltage angle V_{a,i}. "
        "Given N load scenarios solved by NR, a GNN f_\u03b8 is trained "
        "to approximate the power flow solution, with mean squared error "
        "(MSE) loss:")

    # Formula 1: MSE loss (use text fallback for complex expression)
    add_inline_formula_para(doc,
        "L_MSE = (1/N|V|) \u03a3_n \u03a3_i \u2016 f_\u03b8(G_n)_i \u2212 y_i^(n) \u2016\u00b2",
        1)

    add_body(doc,
        "where y_i^(n)=[V_{m,i}, V_{a,i}]\u1d40 is the NR solution for "
        "bus i in scenario n.")

    add_heading_unnumbered(doc, "Virtual slack edges.")
    add_body(doc,
        "In a radial distribution feeder, every bus is electrically "
        "connected to the substation (slack bus) through a unique path. "
        "A 4-layer GNN can only aggregate information from nodes within "
        "four hops, but buses at the end of long LV feeders may be 56 "
        "hops from the slack bus, meaning the voltage reference "
        "information never reaches them during message passing. This is "
        "addressed by adding bidirectional edges between the slack bus s "
        "and every other node:")

    add_inline_formula_para(doc,
        "E' = E \u222a {(s,i) : i \u2208 V\\{s}} \u222a {(i,s) : i \u2208 V\\{s}}",
        2)

    add_mixed(doc, [
        ("Proposition 1 (Diameter bound). ", True, False),
        ("Let G=(V,E) be any connected graph with a distinguished vertex "
         "s\u2208V. The augmented graph G'=(V,E') defined by (2) "
         "satisfies diam(G') \u2264 2.", False, False),
    ], after=2, indent=FIRST_INDENT)

    add_mixed(doc, [
        ("Proof. ", False, True),
        ("For any u, v \u2208 V\\{s} with u\u2260v, the path "
         "u\u2192s\u2192v has length 2 in G' by construction. Paths "
         "involving s itself have length at most 1. \u25a1", False, False),
    ], after=4, indent=FIRST_INDENT)

    add_mixed(doc, [
        ("Corollary 1 (One-hop reachability). ", True, False),
        ("A single message-passing layer on G' suffices for information "
         "to propagate from the slack bus s to every other node.",
         False, False),
    ], after=4, indent=FIRST_INDENT)

    add_body(doc,
        "Proposition 1 is the structural guarantee behind the approach: "
        "whereas a d-layer GNN on the original radial feeder G requires "
        "d\u2265diam(G) to propagate the slack reference to every bus "
        "(up to 56 layers for LV grids in our benchmark), on G' a single "
        "layer always suffices. The added edges carry zero-impedance "
        "features, distinguishing them from physical branches. This "
        "approach differs from the generic virtual node technique [20] "
        "in that it reuses an existing physically meaningful node rather "
        "than introducing a new artificial one. The number of additional "
        "edges is 2(|V|\u22121), adding at most 286 edges for our grids "
        "(15\u2013144 buses).")

    add_heading_unnumbered(doc, "Random walk positional encodings.")
    add_body(doc,
        "Each node is augmented with random walk positional encodings "
        "(RW-PE) [15]. For a graph with adjacency matrix A and degree "
        "matrix D, the transition matrix is T=D\u207b\u00b9A. The RW-PE "
        "for node i is:")

    add_inline_formula_para(doc,
        "RW-PE_i = [T_{ii}, T\u00b2_{ii}, \u2026, T^k_{ii}] \u2208 \u211d^k",
        3)

    add_body(doc,
        "where T^j_{ii} is the probability that a random walk starting at "
        "node i returns to i after j steps. We use k=16 steps, which "
        "captures both local structure (short return times for high-"
        "degree nodes) and mesoscale topology (longer cycles). RW-PE is "
        "preferred over Laplacian eigenvector PE because (1) it does not "
        "suffer from eigenvector sign ambiguity [14], and (2) Laplacian "
        "PE with k=8 eigenvectors yielded substantially worse results "
        "(MAE 3.02\u00d710\u207b\u00b3 p.u. vs. 0.80\u00d710\u207b\u00b3 "
        "p.u. for RW-PE), likely because the low-frequency eigenvectors "
        "of radial graphs are highly degenerate.")

    add_heading_unnumbered(doc, "Residual GraphSAGE.")
    add_body(doc,
        "To expand the receptive field beyond four hops while avoiding "
        "oversmoothing, a pre-normalization residual architecture is "
        "adopted. Each layer l computes:")

    add_inline_formula_para(doc,
        "h_i^{(l+1)} = h_i^{(l)} + SAGE(LN(h_i^{(l)}), {LN(h_j^{(l)}) : "
        "j \u2208 N(i)})",
        4)

    add_body(doc,
        "where LN denotes layer normalization and SAGE is the GraphSAGE "
        "aggregation (mean aggregator followed by a linear projection) "
        "[22]. The residual connection ensures that even with 8 layers, "
        "gradient flow is maintained and node features do not collapse. "
        "Residual connections, JKNet [18], and DropEdge [19] were "
        "compared in a pilot study (Table 3).")

    add_heading_unnumbered(doc, "Physics-informed loss (attempted).")
    add_body(doc,
        "A voltage-smoothness regularization term penalizing large "
        "voltage differences between connected buses was tested:")

    add_inline_formula_para(doc,
        "L_smooth = (1/|E|) \u03a3_{(i,j)\u2208E} (V\u0302_{m,i} \u2212 "
        "V\u0302_{m,j})\u00b2",
        5)

    add_body(doc,
        "combined with a bounds penalty:")

    add_inline_formula_para(doc,
        "L_bounds = (1/|V|) \u03a3_i max(0, |V\u0302_{m,i} \u2212 1.0| \u2212 0.05)\u00b2",
        6)

    add_body(doc,
        "Total loss was L = L_MSE + \u03bb (L_smooth + L_bounds) with "
        "\u03bb \u2208 {10, 100, 1000}. None of these configurations "
        "improved upon the MSE-only baseline; see Results.")

    # =================================================================
    # 4. EXPERIMENTAL SETUP
    # =================================================================
    add_heading_numbered(doc, "Experimental setup.")

    add_body(doc,
        "Ten distribution grids from the SimBench benchmark [23] are "
        "used, comprising four MV grids (20 kV) and six LV grids "
        "(0.4 kV). Table 1 summarizes their structural properties. LV "
        "grids exhibit substantially larger diameters (up to 56 hops "
        "for LV_rural2) and higher bridge fractions compared with MV "
        "grids, confirming the structural basis for the MV/LV "
        "performance gap.")

    add_table(doc,
        headers=["Grid", "Buses", "Edges", "Diam.", "R/X ratio",
                 "Bridge frac."],
        rows=[
            ["MV_rural", "97", "96", "22", "1.02", "0.96"],
            ["MV_semiurb", "116", "119", "14", "0.98", "0.89"],
            ["MV_urban", "144", "147", "12", "0.85", "0.88"],
            ["MV_comm", "89", "88", "20", "1.05", "0.97"],
            ["LV_rural1", "15", "14", "13", "5.21", "1.00"],
            ["LV_rural2", "96", "95", "56", "4.87", "1.00"],
            ["LV_rural3", "60", "59", "38", "4.92", "1.00"],
            ["LV_semiurb4", "44", "43", "28", "3.45", "1.00"],
            ["LV_semiurb5", "55", "54", "32", "3.68", "1.00"],
            ["LV_urban6", "40", "41", "18", "2.95", "0.95"],
        ],
        caption_num=1,
        caption_text="Structural properties of SimBench distribution grids.")

    add_body(doc,
        "For each grid, 2,000 load scenarios are generated by sampling "
        "multiplicative factors uniformly from [0.7, 1.3] applied to the "
        "nominal active and reactive power loads, representing \u00b130% "
        "load variation. Each scenario is solved using the pandapower NR "
        "solver [24]. Four critical corrections are applied to SimBench "
        "models: transformer shift degrees set to zero, continuous bus "
        "reindexing, NaN voltage-dependent load parameters replaced with "
        "zero, and all switches removed. These corrections are essential "
        "for NR convergence on the full set of grids.")

    add_body(doc,
        "Data are split into training (70%), validation (15%), and test "
        "(15%) sets. Input features and targets are Z-score normalized "
        "per grid. All models use AdamW (learning rate 10\u207b\u00b3, "
        "weight decay 10\u207b\u2074) with OneCycleLR scheduling over "
        "200 epochs, early stopping with patience 25, hidden dimension "
        "64. Results are averaged over three random seeds (42, 123, "
        "456). All experiments run on CPU (Apple M1 Pro) to ensure fair "
        "timing comparisons between GNN inference and the NR solver.")

    # =================================================================
    # 5. RESULTS
    # =================================================================
    add_heading_numbered(doc, "Results.")

    add_heading_unnumbered(doc, "Ablation study.")
    add_body(doc,
        "Table 2 summarizes the individual contributions of each proposed "
        "component relative to the 4-layer GraphSAGE baseline.")

    add_table(doc,
        headers=["Method", "MAE V_m (\u00d710\u207b\u00b3 p.u.)",
                 "LV/MV gap", "Improv. (%)"],
        rows=[
            ["Baseline (GraphSAGE d=4)", "0.88 \u00b1 0.25", "1.41\u00d7", "\u2014"],
            ["E1: Physics loss (\u03bb=10)", "0.88 \u00b1 0.25", "1.41\u00d7", "~0"],
            ["E1: Physics loss (\u03bb=100)", "0.88 \u00b1 0.25", "1.41\u00d7", "~0"],
            ["E1: Physics loss (\u03bb=1000)", "0.88 \u00b1 0.25", "1.41\u00d7", "~0"],
            ["E2: RW-PE (k=16)", "0.80 \u00b1 0.21", "1.38\u00d7", "9"],
            ["E2: Dist-from-slack", "0.81 \u00b1 0.22", "1.42\u00d7", "8"],
            ["E2: Laplacian PE (k=8)", "3.02 \u00b1 1.89", "1.41\u00d7", "\u2212244"],
            ["E4: Virtual slack edges", "0.58", "1.20\u00d7", "34"],
        ],
        caption_num=2,
        caption_text="Ablation study: individual components vs. baseline.")

    add_body(doc,
        "Physics-informed loss (E1). Across all three values of \u03bb, "
        "the physics-informed loss produced no measurable improvement. "
        "Averaged over ten grids and three seeds, the mean MAE change "
        "relative to baseline was \u22120.004\u00d710\u207b\u00b3, "
        "\u22120.006\u00d710\u207b\u00b3, and \u22120.004\u00d710\u207b\u00b3 "
        "p.u. for \u03bb\u2208{10,100,1000} respectively, with the "
        "largest single-grid deviation only 0.03\u00d710\u207b\u00b3 "
        "p.u. \u2014 well within seed-to-seed noise and two orders of "
        "magnitude smaller than the 0.44\u00d710\u207b\u00b3 p.u. "
        "improvement achieved by the combined model.")

    add_body(doc,
        "Positional encodings (E2). Random walk PE (k=16) achieved a "
        "modest 9% improvement in overall MAE and slightly reduced the "
        "LV/MV gap from 1.41\u00d7 to 1.38\u00d7. Distance-from-slack "
        "PE performed comparably (8% improvement) but did not reduce the "
        "gap. Laplacian PE degraded performance severely, likely due to "
        "eigenvector sign ambiguity and the near-degenerate spectrum of "
        "tree-structured graphs.")

    add_body(doc,
        "Virtual slack edges (E4). This was the most effective single "
        "component, reducing overall MAE by 34% and compressing the "
        "LV/MV gap from 1.41\u00d7 to 1.20\u00d7. The MV average dropped "
        "from 0.70\u00d710\u207b\u00b3 p.u. to 0.52\u00d710\u207b\u00b3 "
        "p.u., while the LV average dropped from 0.99\u00d710\u207b\u00b3 "
        "p.u. to 0.63\u00d710\u207b\u00b3 p.u. The improvement is "
        "disproportionately larger for LV grids, confirming that the "
        "virtual edges specifically address the long-diameter bottleneck.")

    add_heading_unnumbered(doc, "Depth study.")
    add_body(doc,
        "A pilot study was conducted on two grids (MV_rural and "
        "LV_rural2) to evaluate three anti-oversmoothing techniques at "
        "depths 4, 8, 16, and 32 (Table 3). Residual connections at d=8 "
        "achieved the best MAE (0.91\u00d710\u207b\u00b3), while JKNet "
        "provided the most stable performance across depths (least "
        "degradation at d=32). DropEdge was consistently inferior. At "
        "d=32, both Residual and DropEdge suffered severe oversmoothing, "
        "whereas JKNet maintained reasonable accuracy by selectively "
        "combining layer representations. Residual connections with d=8 "
        "were therefore selected for the combined model.")

    add_table(doc,
        headers=["Technique", "d=4", "d=8", "d=16", "d=32"],
        rows=[
            ["Residual", "1.02", "0.91", "0.93", "4.30"],
            ["JKNet", "1.03", "0.93", "0.95", "1.16"],
            ["DropEdge", "1.19", "1.02", "1.11", "4.77"],
        ],
        caption_num=3,
        caption_text="Depth study: MAE V_m (\u00d710\u207b\u00b3 p.u.) "
                     "by technique and depth.")

    add_heading_unnumbered(doc, "Combined model.")
    add_body(doc,
        "Table 4 presents the main result: the combined model (Residual "
        "GraphSAGE d=8 + RW-PE k=16 + virtual slack edges) compared with "
        "the baseline across all ten grids. The combined model reduces "
        "MAE V_m on every grid without exception. The largest absolute "
        "improvement is on LV_rural2 (1.37\u21920.75, a 45% reduction), "
        "which has the largest diameter (56 hops) and was the single "
        "worst-performing grid for the baseline. Voltage angle MAE "
        "improves even more dramatically, with reductions of 47\u201374% "
        "across grids. A Wilcoxon signed-rank test across all ten grids "
        "confirms that the improvement is statistically significant "
        "(p=0.002), with a mean improvement of 43.8%\u00b110.5%.")

    add_table(doc,
        headers=["Grid", "Baseline MAE V_m", "Combined MAE V_m",
                 "Baseline MAE V_a", "Combined MAE V_a",
                 "Baseline speedup", "Combined speedup"],
        rows=[
            ["MV_rural", "0.75 \u00b1 0.02", "0.60 \u00b1 0.04",
             "0.096", "0.029", "25.1\u00d7", "6.1\u00d7"],
            ["MV_semiurb", "0.72 \u00b1 0.02", "0.44 \u00b1 0.02",
             "0.089", "0.026", "23.5\u00d7", "3.4\u00d7"],
            ["MV_urban", "0.60 \u00b1 0.03", "0.30 \u00b1 0.01",
             "0.062", "0.033", "25.1\u00d7", "1.5\u00d7"],
            ["MV_comm", "0.75 \u00b1 0.04", "0.42 \u00b1 0.02",
             "0.097", "0.031", "25.1\u00d7", "2.7\u00d7"],
            ["LV_rural1", "0.70 \u00b1 0.05", "0.34 \u00b1 0.15",
             "0.087", "0.037", "34.0\u00d7", "2.6\u00d7"],
            ["LV_rural2", "1.37 \u00b1 0.05", "0.75 \u00b1 0.02",
             "0.117", "0.027", "26.5\u00d7", "2.2\u00d7"],
            ["LV_rural3", "1.11 \u00b1 0.04", "0.72 \u00b1 0.01",
             "0.077", "0.024", "24.0\u00d7", "2.5\u00d7"],
            ["LV_semiurb4", "0.95 \u00b1 0.04", "0.40 \u00b1 0.05",
             "0.052", "0.014", "29.2\u00d7", "3.1\u00d7"],
            ["LV_semiurb5", "1.15 \u00b1 0.02", "0.62 \u00b1 0.02",
             "0.069", "0.022", "24.3\u00d7", "4.1\u00d7"],
            ["LV_urban6", "0.68 \u00b1 0.01", "0.34 \u00b1 0.01",
             "0.047", "0.012", "29.4\u00d7", "2.8\u00d7"],
            ["MV avg", "0.70", "0.44", "0.086", "0.030",
             "24.7\u00d7", "3.4\u00d7"],
            ["LV avg", "0.99", "0.53", "0.075", "0.023",
             "27.9\u00d7", "2.9\u00d7"],
        ],
        caption_num=4,
        caption_text="Per-grid comparison: baseline GraphSAGE vs. combined "
                     "model. MAE V_m in \u00d710\u207b\u00b3 p.u.; MAE V_a "
                     "in degrees; speedup relative to Newton\u2013Raphson.",
        bold_rows={10, 11})

    add_heading_unnumbered(doc, "MV/LV gap analysis.")
    add_body(doc,
        "Fig. 1 illustrates the distribution of MAE V_m for MV and LV "
        "grids before and after applying the combined model. For the "
        "baseline, the LV/MV ratio is 1.41\u00d7, and the difference is "
        "statistically significant at the per-seed level (Mann\u2013"
        "Whitney U test on 12 MV vs. 18 LV observations, p=0.010). "
        "After applying the combined model, the ratio drops to 1.21\u00d7, "
        "and the gap becomes statistically non-significant (p=0.182) on "
        "the SimBench benchmark used here.")

    add_figure(doc, 1,
        "Distribution of MAE V_m for MV and LV grids. Left: baseline "
        "GraphSAGE shows a significant gap (Mann\u2013Whitney p=0.010). "
        "Right: the combined model reduces the gap to non-significance "
        "(p=0.182).",
        os.path.join(fig_dir, "fig1_mvlv_gap_boxplot.png"), width_cm=14.0)

    add_heading_unnumbered(doc, "Structural correlates of the gap.")
    add_body(doc,
        "To move from attribution to evidence, three structural "
        "properties of each grid \u2014 diameter, bridge fraction, and "
        "R/X ratio \u2014 are correlated with baseline MAE and with the "
        "improvement delivered by the combined model (Table 5).")

    add_table(doc,
        headers=["Property", "vs. Baseline MAE", "vs. Combined \u0394MAE"],
        rows=[
            ["Diameter",
             "\u03c1=+0.95, p<10\u207b\u2074",
             "\u03c1=+0.66, p=0.038"],
            ["Bridge fraction",
             "\u03c1=+0.74, p=0.014",
             "\u03c1=+0.82, p=0.004"],
            ["R/X ratio",
             "\u03c1=+0.50, p=0.138",
             "\u03c1=+0.75, p=0.013"],
        ],
        caption_num=5,
        caption_text="Spearman rank correlations between grid structure "
                     "and performance (n=10).")

    add_body(doc,
        "Diameter is almost perfectly rank-correlated with baseline "
        "error (\u03c1=+0.95), which constitutes direct evidence for the "
        "receptive-field hypothesis formulated earlier: grids whose "
        "longest path exceeds the receptive field of a 4-layer GNN are "
        "the ones on which the baseline suffers. Bridge fraction "
        "(\u03c1=+0.74) corroborates this, since high-bridge topologies "
        "offer no alternative message-passing paths. The R/X ratio, a "
        "purely electrical property, correlates more weakly with baseline "
        "error once diameter is accounted for. The second column links "
        "structure to the benefit of the method. The absolute improvement "
        "grows with both diameter (\u03c1=+0.66) and bridge fraction "
        "(\u03c1=+0.82): the grids where the receptive-field bottleneck "
        "is most severe are also the ones that gain the most from the "
        "virtual slack edges. This is consistent with the 2-hop diameter "
        "guarantee of Proposition 1 \u2014 its effect scales with how "
        "far the baseline was from the slack reference to begin with.")

    add_heading_unnumbered(doc, "Speed\u2013accuracy trade-off.")
    add_body(doc,
        "The combined model's improved accuracy comes at the cost of "
        "reduced inference speed. The baseline GraphSAGE achieves "
        "24\u201334\u00d7 speedup over NR, while the combined model "
        "achieves 1.5\u20136.1\u00d7 speedup (Fig. 2). This reduction "
        "stems from two factors: (1) the virtual slack edges roughly "
        "double the edge count, increasing per-layer computation; and "
        "(2) the deeper architecture (8 vs. 4 layers) doubles the number "
        "of message-passing steps. Nevertheless, the combined model "
        "remains faster than NR on all grids (minimum 1.5\u00d7 on "
        "MV_urban), providing a net computational benefit. For accuracy-"
        "critical applications \u2014 voltage violation detection or "
        "hosting capacity analysis \u2014 the combined model is "
        "preferable; the baseline remains better when maximum throughput "
        "matters more than accuracy.")

    add_figure(doc, 2,
        "Speed\u2013accuracy Pareto frontier. The combined model "
        "(triangles) achieves lower MAE at the cost of reduced speedup. "
        "The baseline (circles) offers higher speedup with moderate "
        "accuracy.",
        os.path.join(fig_dir, "fig2_pareto_speedup_accuracy.png"),
        width_cm=14.0)

    add_figure(doc, 3,
        "Per-grid MAE V_m comparison between baseline GraphSAGE and the "
        "combined model. Error bars indicate \u00b11 standard deviation "
        "over three seeds.",
        os.path.join(fig_dir, "fig3_per_grid_comparison.png"), width_cm=14.0)

    # =================================================================
    # 6. DISCUSSION
    # =================================================================
    add_heading_numbered(doc, "Discussion.")

    add_heading_unnumbered(doc, "Why virtual slack edges dominate.")
    add_body(doc,
        "The virtual slack edges produced the largest single improvement "
        "(34% MAE reduction), substantially outperforming positional "
        "encodings (9%) and deeper architectures alone (11% on the pilot "
        "grids). This dominance has a clear structural explanation: in a "
        "radial feeder, the slack bus is the voltage reference, and "
        "voltage drops accumulate along the unique path from slack to "
        "each load bus. By providing a direct 2-hop connection from "
        "every bus to the slack, the virtual edges allow the GNN to "
        "\u201csee\u201d the reference voltage in a single message-"
        "passing step, regardless of the electrical distance. This is "
        "more effective than simply deepening the network because a "
        "deep GNN still processes information sequentially through "
        "intermediate nodes, with potential information loss at each "
        "layer. The virtual edges provide a shortcut that bypasses this "
        "sequential bottleneck entirely.")

    add_heading_unnumbered(doc, "Why physics-informed loss fails.")
    add_body(doc,
        "Physics-informed approaches have shown promise for optimal "
        "power flow [7, 8], yet our voltage-smoothness loss had no "
        "effect. A scale analysis explains the result. In SimBench "
        "distribution grids, adjacent buses typically differ in voltage "
        "magnitude by |V_i \u2212 V_j| \u223c 10\u207b\u00b3 p.u. under "
        "\u00b130% load variation, so L_smooth on the ground-truth NR "
        "solutions is already of order 10\u207b\u2076. Because we train "
        "on Z-score normalized targets, the MSE term operates in the "
        "normalized scale while L_smooth remains anchored to the "
        "physical scale; the \u03bb values we swept (10, 100, 1000) "
        "span three orders of magnitude and still fail to produce a "
        "detectable effect, confirming that the regularization is "
        "dominated by, and redundant with, the MSE signal on normalized "
        "NR solutions.")

    add_heading_unnumbered(doc, "Limitations.")
    add_body(doc,
        "Our grids have at most 144 buses, which is small relative to "
        "real-world feeders that may contain thousands of nodes. The "
        "virtual slack edge approach adds O(|V|) edges, which may become "
        "computationally significant for very large grids. The speedup "
        "reduction from 25\u00d7 to 2\u20136\u00d7 limits the practical "
        "advantage for latency-sensitive applications. Our evaluation is "
        "confined to the SimBench benchmark: all reported statistics "
        "\u2014 including the structural correlations \u2014 describe "
        "behavior on this particular family of grids and should not be "
        "read as generalizable to real distribution networks without "
        "further validation. All grids are radial; the 2-hop guarantee "
        "of Proposition 1 still holds for meshed topologies, but the "
        "receptive-field pressure that makes the method especially "
        "useful may be weaker there, and we do not test this empirically. "
        "We do not consider distributed generation or stochastic "
        "renewable profiles, which would add variability beyond the "
        "\u00b130% load perturbation used here. We do not compare "
        "against recent physics-informed alternatives such as PINCO [7] "
        "or line graph formulations [8]; such comparisons are left for "
        "future work.")

    add_heading_unnumbered(doc, "Practical implications.")
    add_body(doc,
        "For distribution system operators running Monte Carlo studies "
        "for hosting capacity assessment or probabilistic voltage "
        "analysis, a model that is 2\u20136\u00d7 faster than NR while "
        "maintaining sub-10\u207b\u00b3 p.u. voltage-magnitude accuracy "
        "offers tangible benefits. The reduction of the MV/LV "
        "performance gap to statistical non-significance on SimBench is "
        "particularly valuable, since LV grids are precisely where the "
        "computational burden of repeated power flow is greatest due to "
        "their larger number and the need for detailed analysis of "
        "voltage-sensitive loads; whether the same holds on utility-"
        "scale feeders with different topological statistics is an "
        "empirical question that we do not answer here.")

    # =================================================================
    # 7. CONCLUSIONS
    # =================================================================
    add_heading_numbered(doc, "Conclusions.")

    add_body(doc,
        "1. Three targeted techniques for GNN-based AC power flow on "
        "radial distribution networks were proposed and evaluated: "
        "virtual slack edges, random walk positional encodings, and "
        "residual GraphSAGE with eight layers. Each was tested "
        "individually and in combination across ten SimBench distribution "
        "grids (four MV, six LV) with 2,000 load scenarios and three "
        "random seeds per grid. The combined model reduces voltage-"
        "magnitude MAE by 43.8% on average (Wilcoxon p=0.002) relative "
        "to a standard 4-layer GraphSAGE baseline, with improvements on "
        "every grid without exception.")

    add_body(doc,
        "2. The MV/LV performance gap measured on this benchmark is "
        "reduced from 1.41\u00d7 (p=0.010) to 1.21\u00d7 (p=0.182), "
        "rendering it statistically non-significant. The receptive-field "
        "hypothesis that motivated the design is supported by a strong "
        "rank correlation between graph diameter and baseline MAE "
        "(\u03c1=+0.95, p<10\u207b\u2074), and the absolute improvement "
        "grows with diameter and bridge fraction, consistent with the "
        "2-hop diameter guarantee established in Proposition 1. Virtual "
        "slack edges are the single most impactful component, providing "
        "a 34% MAE reduction.")

    add_body(doc,
        "3. A physics-informed voltage-smoothness loss was tested with "
        "\u03bb \u2208 {10, 100, 1000} and produced no measurable "
        "improvement over the MSE-only baseline; a scale analysis "
        "explains why the regularization is redundant on Z-score "
        "normalized NR solutions. This is a negative result of practical "
        "value: for power flow approximation, unlike optimal power flow, "
        "additional physics constraints on the loss function should not "
        "be expected to help unless the training targets themselves fail "
        "to satisfy the physics. Future work should investigate "
        "scalability to grids with 1,000 or more buses (where the "
        "quadratic cost of virtual edges may necessitate sparse "
        "approximations), transfer learning across grid topologies, and "
        "the incorporation of time-varying load and generation profiles.")

    # =================================================================
    # Acknowledgments
    # =================================================================
    add_heading_unnumbered(doc, "Acknowledgments.")
    add_body(doc,
        "The authors thank the SimBench consortium for providing the "
        "open-access distribution grid benchmark data.")

    # =================================================================
    # Conflict of interest
    # =================================================================
    add_heading_unnumbered(doc, "Conflict of interest.")
    add_body(doc,
        "The authors declare that they have no known competing financial "
        "interests or personal relationships that could have appeared to "
        "influence the work reported in this paper.")

    # =================================================================
    # Data availability
    # =================================================================
    add_heading_unnumbered(doc, "Data availability.")
    add_body(doc,
        "The SimBench grid models used in this study are publicly "
        "available at https://simbench.de/en/. All source code and trained "
        "model checkpoints are released under the MIT license at "
        "https://github.com/dvoitekh/gnn-power-flow-distribution and will "
        "be archived on Zenodo with a persistent DOI upon acceptance.")

    # =================================================================
    # AI disclosure
    # =================================================================
    add_heading_unnumbered(doc,
        "Declaration of generative AI and AI-assisted technologies.")
    add_body(doc,
        "During the preparation of this work, the authors used Claude "
        "(Anthropic) for code development assistance. The authors "
        "reviewed and edited all content and take full responsibility "
        "for the publication.")

    # =================================================================
    # CRediT
    # =================================================================
    add_heading_unnumbered(doc, "Authors' contributions.")
    add_body(doc,
        "Dmytro Voitekh: conceptualization, methodology, software, "
        "validation, formal analysis, investigation, data curation, "
        "writing \u2014 original draft, visualization. "
        "Anatolii Tymoshenko: supervision, writing \u2014 review and "
        "editing.")

    # =================================================================
    # REFERENCES
    # =================================================================
    add_heading_unnumbered(doc, "REFERENCES")

    references = [
        # 1
        "Wang Y., Chen Q., Hong T. et al. Review of smart meter data "
        "analytics: applications, methodologies, and challenges. IEEE "
        "Transactions on Smart Grid, 2019, vol. 10, no. 3, pp. 3125\u20133148. "
        "DOI: https://doi.org/10.1109/TSG.2017.2778801.",
        # 2
        "Capitanescu F. Critical review of recent advances and further "
        "developments needed in AC optimal power flow. Electric Power "
        "Systems Research, 2016, vol. 136, pp. 57\u201368. DOI: "
        "https://doi.org/10.1016/j.epsr.2016.02.008.",
        # 3
        "Kipf T.N., Welling M. Semi-supervised classification with graph "
        "convolutional networks. Proceedings of the 5th International "
        "Conference on Learning Representations (ICLR), 2017, 14 p. DOI: "
        "https://doi.org/10.48550/arXiv.1609.02907.",
        # 4
        "Donon B., Cl\u00e9ment R., Donnot B. et al. Neural networks for "
        "power flow: graph neural solver. Electric Power Systems Research, "
        "2020, vol. 189, art. no. 106547. DOI: "
        "https://doi.org/10.1016/j.epsr.2020.106547.",
        # 5
        "Lin N., Orfanoudakis S., Ordonez Cardenas N. et al. PowerFlowNet: "
        "power flow approximation using message passing graph neural "
        "networks. International Journal of Electrical Power & Energy "
        "Systems, 2024, vol. 160, art. no. 110112. DOI: "
        "https://doi.org/10.1016/j.ijepes.2024.110112.",
        # 6
        "Ringsquandl M., Sellami H., Hildebrandt M. et al. Power to the "
        "relational inductive bias: graph neural networks in electrical "
        "power grids. Proceedings of the 30th ACM International Conference "
        "on Information & Knowledge Management (CIKM), 2021, pp. 1538\u20131547. "
        "DOI: https://doi.org/10.1145/3459637.3482464.",
        # 7
        "Varbella A., Briens D., Gjorgiev B. et al. Physics-informed GNN "
        "for non-linear constrained optimization: PINCO, a solver for the "
        "AC-optimal power flow. arXiv preprint arXiv:2410.04818, 2024. "
        "DOI: https://doi.org/10.48550/arXiv.2410.04818.",
        # 8
        "Zhang H.-F., Lu X.-L., Ding X., Zhang X.-M. Physics-informed line "
        "graph neural network for power flow calculation. Chaos, 2024, "
        "vol. 34, no. 11, art. no. 113123. DOI: "
        "https://doi.org/10.1063/5.0235355.",
        # 9
        "Varbella A., Amara K., Gjorgiev B. et al. PowerGraph: a power grid "
        "benchmark dataset for graph neural networks. Advances in Neural "
        "Information Processing Systems 37 (NeurIPS) Datasets and "
        "Benchmarks Track, 2024. DOI: "
        "https://doi.org/10.48550/arXiv.2402.02827.",
        # 10
        "Ghamizi S., Cao J., Ma A., Rodriguez P. PowerFlowMultiNet: "
        "multigraph neural networks for unbalanced three-phase distribution "
        "systems. IEEE Transactions on Power Systems, 2024. DOI: "
        "https://doi.org/10.1109/TPWRS.2024.3465088.",
        # 11
        "Suri D., Mangal M. PowerGNN: a topology-aware graph neural "
        "network for electricity grids. arXiv preprint arXiv:2503.22721, "
        "2025. DOI: https://doi.org/10.48550/arXiv.2503.22721.",
        # 12
        "Okoyomon E., Goebel C. A framework for assessing the "
        "generalizability of GNN-based AC power flow models. Proceedings "
        "of the 16th ACM International Conference on Future and "
        "Sustainable Energy Systems (e-Energy), 2025. DOI: "
        "https://doi.org/10.1145/3679240.3734610.",
        # 13
        "Xu K., Hu W., Leskovec J., Jegelka S. How powerful are graph "
        "neural networks? Proceedings of the 7th International Conference "
        "on Learning Representations (ICLR), 2019, 17 p. DOI: "
        "https://doi.org/10.48550/arXiv.1810.00826.",
        # 14
        "Dwivedi V.P., Joshi C.K., Luu A.T. et al. Benchmarking graph "
        "neural networks. Journal of Machine Learning Research, 2023, "
        "vol. 24, no. 43, pp. 1\u201348. DOI: "
        "https://doi.org/10.48550/arXiv.2003.00982.",
        # 15
        "Li P., Wang Y., Wang H., Leskovec J. Distance encoding: design "
        "provably more powerful neural networks for graph representation "
        "learning. Advances in Neural Information Processing Systems 33 "
        "(NeurIPS), 2020, pp. 4465\u20134478. DOI: "
        "https://doi.org/10.48550/arXiv.2009.00142.",
        # 16
        "Li Q., Han Z., Wu X.-M. Deeper insights into graph convolutional "
        "networks for semi-supervised learning. Proceedings of the 32nd "
        "AAAI Conference on Artificial Intelligence, 2018, pp. 3538\u20133545. "
        "DOI: https://doi.org/10.1609/aaai.v32i1.11604.",
        # 17
        "He K., Zhang X., Ren S., Sun J. Deep residual learning for image "
        "recognition. Proceedings of the IEEE Conference on Computer "
        "Vision and Pattern Recognition (CVPR), 2016, pp. 770\u2013778. "
        "DOI: https://doi.org/10.1109/CVPR.2016.90.",
        # 18
        "Xu K., Li C., Tian Y. et al. Representation learning on graphs "
        "with jumping knowledge networks. Proceedings of the 35th "
        "International Conference on Machine Learning (ICML), 2018, "
        "pp. 5453\u20135462. DOI: "
        "https://doi.org/10.48550/arXiv.1806.03536.",
        # 19
        "Rong Y., Huang W., Xu T., Huang J. DropEdge: towards deep graph "
        "convolutional networks on node classification. Proceedings of "
        "the 8th International Conference on Learning Representations "
        "(ICLR), 2020, 17 p. DOI: "
        "https://doi.org/10.48550/arXiv.1907.10903.",
        # 20
        "Gilmer J., Schoenholz S.S., Riley P.F. et al. Neural message "
        "passing for quantum chemistry. Proceedings of the 34th "
        "International Conference on Machine Learning (ICML), 2017, "
        "pp. 1263\u20131272. DOI: "
        "https://doi.org/10.48550/arXiv.1704.01212.",
        # 21
        "Hu W., Fey M., Zitnik M. et al. Open graph benchmark: datasets "
        "for machine learning on graphs. Advances in Neural Information "
        "Processing Systems 33 (NeurIPS), 2020, pp. 22118\u201322133. "
        "DOI: https://doi.org/10.48550/arXiv.2005.00687.",
        # 22
        "Hamilton W.L., Ying R., Leskovec J. Inductive representation "
        "learning on large graphs. Advances in Neural Information "
        "Processing Systems 30 (NeurIPS), 2017, pp. 1025\u20131035. DOI: "
        "https://doi.org/10.48550/arXiv.1706.02216.",
        # 23
        "Meinecke S., Sarajli\u0107 D., Drauz S.R. et al. SimBench \u2014 "
        "a benchmark dataset of electric power systems to compare "
        "innovative solutions based on power flow analysis. Energies, "
        "2020, vol. 13, no. 12, art. no. 3290. DOI: "
        "https://doi.org/10.3390/en13123290.",
        # 24
        "Thurner L., Scheidler A., Sch\u00e4fer F. et al. pandapower \u2014 "
        "an open-source Python tool for convenient modeling, analysis, "
        "and optimization of electric power systems. IEEE Transactions on "
        "Power Systems, 2018, vol. 33, no. 6, pp. 6510\u20136521. DOI: "
        "https://doi.org/10.1109/TPWRS.2018.2829021.",
        # 25
        "Hu X., Yang J., Gao Y. et al. Adaptive power flow analysis for "
        "power system operation based on graph deep learning. "
        "International Journal of Electrical Power & Energy Systems, "
        "2024, vol. 160, art. no. 110166. DOI: "
        "https://doi.org/10.1016/j.ijepes.2024.110166.",
        # 26
        "Ngo Q.-H., Nguyen B.L.H., Vu T., Zhang J. Physics-informed "
        "graphical neural network for power system state estimation. "
        "Applied Energy, 2024, vol. 358, art. no. 122602. DOI: "
        "https://doi.org/10.1016/j.apenergy.2023.122602.",
        # 27
        "Beinert D., Holzh\u00fcter C., Thomas J.M., Vogt S. Power flow "
        "forecasts at transmission grid nodes using graph neural "
        "networks. Energy and AI, 2023, vol. 14, art. no. 100262. DOI: "
        "https://doi.org/10.1016/j.egyai.2023.100262.",
        # 28
        "Madbhavi R., Natarajan B., Srinivasan B. Graph neural network-"
        "based distribution system state estimators. IEEE Transactions on "
        "Industrial Informatics, 2023, vol. 19, no. 11, pp. 11383\u201311393. "
        "DOI: https://doi.org/10.1109/TII.2023.3248082.",
        # 29
        "Lopez-Garcia T.B., Dom\u00ednguez-Navarro J.A. Power flow "
        "analysis via typed graph neural networks. Engineering "
        "Applications of Artificial Intelligence, 2023, vol. 117, art. "
        "no. 105567. DOI: "
        "https://doi.org/10.1016/j.engappai.2022.105567.",
        # 30
        "Gao M., Yu J., Yang Z., Zhao J. Physics embedded graph "
        "convolution neural network for power flow calculation "
        "considering uncertain injections and topology. IEEE Transactions "
        "on Neural Networks and Learning Systems, 2024. DOI: "
        "https://doi.org/10.1109/TNNLS.2023.3287028.",
        # 31
        "Gao M., Yu J., Yang Z. A physics-guided graph convolution neural "
        "network for optimal power flow. IEEE Transactions on Power "
        "Systems, 2024, vol. 39, no. 1, pp. 380\u2013390. DOI: "
        "https://doi.org/10.1109/TPWRS.2023.3238377.",
        # 32
        "Deihim A., Apostolopoulou D., Alonso E. Initial estimate of AC "
        "optimal power flow with graph neural networks. Electric Power "
        "Systems Research, 2024, vol. 234, art. no. 110782. DOI: "
        "https://doi.org/10.1016/j.epsr.2024.110782.",
        # 33
        "Liao W., Bak-Jensen B., Pillai J.R., Wang Y., Wang Y. A review "
        "of graph neural networks and their applications in power "
        "systems. Journal of Modern Power Systems and Clean Energy, "
        "2022, vol. 10, no. 2, pp. 345\u2013360. DOI: "
        "https://doi.org/10.35833/MPCE.2021.000058.",
        # 34
        "Chen Y., Liu Y., Zhao J., Qiu G., Yin H., Li Z. Physical-assisted "
        "multi-agent graph reinforcement learning enabled fast voltage "
        "regulation for PV-rich active distribution network. Applied "
        "Energy, 2023, vol. 351, art. no. 121743. DOI: "
        "https://doi.org/10.1016/j.apenergy.2023.121743.",
        # 35
        "Pei Y., Zhao J., Yao Y., Ding F. Multi-task reinforcement "
        "learning for distribution system voltage control with topology "
        "changes. IEEE Transactions on Smart Grid, 2023, vol. 14, no. 3, "
        "pp. 2481\u20132484. DOI: "
        "https://doi.org/10.1109/TSG.2022.3233766.",
        # 36
        "Zhang Y., Karve P.M., Mahadevan S. Graph neural networks for "
        "power grid operational risk assessment under evolving unit "
        "commitment. Applied Energy, 2025, vol. 380, art. no. 124793. "
        "DOI: https://doi.org/10.1016/j.apenergy.2024.124793.",
    ]
    for i, ref in enumerate(references, start=1):
        add_reference(doc, i, ref)

    # =================================================================
    # Author info block at the end
    # =================================================================
    doc.add_paragraph().paragraph_format.line_spacing = 1.0

    add_para(doc, "Information about all Authors:",
             bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=4)

    author_lines_en = [
        ("D. Voitekh1, PhD Student,", True),
        ("A. Tymoshenko1, PhD, Associate Professor,", True),
        ("1 Department of Computer Engineering, Open International "
         "University of Human Development \u201cUkraine\u201d,", False),
        ("Lvivska Str. 23, Kyiv, 03115, Ukraine,", False),
        ("e-mail: d.voitekh@gmail.com (Corresponding author);", False),
        ("timoshag@i.ua.", False),
        ("ORCID: D. Voitekh https://orcid.org/0009-0003-8997-5495;", False),
        ("A. Tymoshenko https://orcid.org/0000-0003-0954-3186.", False),
    ]
    for text, bold in author_lines_en:
        add_para(doc, text, bold=bold, align=WD_ALIGN_PARAGRAPH.LEFT, after=1)

    add_para(doc, "How to cite this article (by Editorial Board):",
             italic=True, align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=2)
    add_para(doc,
        "Voitekh D., Tymoshenko A. Bridging the MV/LV gap: virtual slack "
        "nodes and positional encodings for GNN-based power flow on "
        "radial distribution networks. Electrical Engineering & "
        "Electromechanics, 202X, no. X, pp. XX\u2013XX. DOI: "
        "https://doi.org/10.20998/2074-272X.202X.X.XX.",
        align=WD_ALIGN_PARAGRAPH.LEFT, after=2)

    out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                            "voitekh_eie_gnn_power_flow.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
    size_kb = os.path.getsize(out_path) // 1024
    print(f"Size:  {size_kb} KB")


if __name__ == "__main__":
    generate()
